import os
import subprocess
from PIL import Image
import fitz  # PyMuPDF
from docx import Document


def convert_document(input_path, target_format, output_dir):
    filename = os.path.basename(input_path)
    name, _ = os.path.splitext(filename)
    output_path = os.path.join(output_dir, f"{name}.{target_format}")
    
    # Map of conversion functions
    conversion_map = {
        ('pdf', 'docx'): _pdf_to_docx,
        ('pdf', 'odt'): _pdf_to_odt,
        ('pdf', 'txt'): _pdf_to_txt,
        ('docx', 'pdf'): _docx_to_pdf,
        ('docx', 'odt'): _docx_to_odt,
        ('docx', 'txt'): _docx_to_txt,
        ('odt', 'pdf'): _odt_to_pdf,
        ('odt', 'docx'): _odt_to_docx,
        ('odt', 'txt'): _odt_to_txt,
        ('txt', 'pdf'): _txt_to_pdf,
        ('txt', 'docx'): _txt_to_docx,
        ('txt', 'odt'): _txt_to_odt,
        ('doc', 'pdf'): _doc_to_pdf,
        ('doc', 'docx'): _doc_to_docx,
        ('doc', 'odt'): _doc_to_odt,
        ('doc', 'txt'): _doc_to_txt,
    }
    
    input_ext = os.path.splitext(input_path)[1][1:].lower()
    conversion_key = (input_ext, target_format.lower())
    
    if conversion_key in conversion_map:
        conversion_map[conversion_key](input_path, output_path)
    else:
        raise ValueError(f"Conversion from {input_ext} to {target_format} not supported")
    
    return output_path

def convert_image(input_path, target_format, output_dir):
    filename = os.path.basename(input_path)
    name, _ = os.path.splitext(filename)
    output_path = os.path.join(output_dir, f"{name}.{target_format}")
    
    with Image.open(input_path) as img:
        if target_format.upper() == 'JPG':
            # Ensure RGB mode for JPG
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
        img.save(output_path, format=target_format.upper())
    
    return output_path

# Document conversion functions
def _pdf_to_docx(input_path, output_path):
    doc = fitz.open(input_path)
    text = ""
    for page in doc:
        text += page.get_text()
    
    document = Document()
    document.add_paragraph(text)
    document.save(output_path)

def _pdf_to_odt(input_path, output_path):
    # Using LibreOffice for better conversion
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'odt', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback to text extraction
        doc = fitz.open(input_path)
        text = ""
        for page in doc:
            text += page.get_text()
        
        with open(output_path, 'w') as f:
            f.write(text)

def _pdf_to_txt(input_path, output_path):
    doc = fitz.open(input_path)
    text = ""
    for page in doc:
        text += page.get_text()
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def _docx_to_pdf(input_path, output_path):
    try:
        # Try using LibreOffice first
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback to using python-docx and reportlab (would need to be implemented)
        raise Exception("PDF conversion requires LibreOffice installed")

def _docx_to_odt(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'odt', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback: extract text and save as ODT
        doc = Document(input_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        with open(output_path, 'w') as f:
            f.write(text)

def _docx_to_txt(input_path, output_path):
    doc = Document(input_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def _odt_to_pdf(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        raise Exception("PDF conversion requires LibreOffice installed")

def _odt_to_docx(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback: extract text and save as DOCX
        # This is a simplified approach - in practice you'd need a proper ODT parser
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)

def _odt_to_txt(input_path, output_path):
    # Simplified text extraction
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def _txt_to_pdf(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                       os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback using reportlab (would need to be implemented)
        raise Exception("PDF conversion requires LibreOffice installed")

def _txt_to_docx(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

def _txt_to_odt(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def _doc_to_pdf(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                        os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback to using python-docx and reportlab (would need to be implemented)
        raise Exception("PDF conversion requires LibreOffice installed")

def _doc_to_docx(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir',
                        os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback: extract text and save as DOCX
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)

def _doc_to_odt(input_path, output_path):
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'odt', '--outdir',
                        os.path.dirname(output_path), input_path], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback: extract text and save as ODT
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

def _doc_to_txt(input_path, output_path):
    # Simplified text extraction
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)